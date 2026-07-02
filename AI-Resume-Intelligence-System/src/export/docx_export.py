from docx import Document

class DOCXExporter:
    @staticmethod
    def export(resume_data: dict, output_path: str) -> bool:
        """Generates an ATS-compliant DOCX resume."""
        try:
            doc = Document()
            
            # Personal info
            info = resume_data.get("personal_info", {})
            doc.add_heading(info.get("full_name", "Resume Candidate"), level=0)
            
            contact_line = f"{info.get('location', '')} | {info.get('phone', '')} | {info.get('email', '')} | {info.get('website', '')}"
            doc.add_paragraph(contact_line)
            
            # Summary
            doc.add_heading("PROFESSIONAL SUMMARY", level=1)
            doc.add_paragraph(resume_data.get("summary", ""))
            
            # Experience
            doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
            for job in resume_data.get("experience", []):
                p = doc.add_paragraph()
                p.add_run(f"{job.get('job_title', '')}").bold = True
                p.add_run(f" – {job.get('company', '')} ({job.get('start_date', '')} - {job.get('end_date', '')})")
                for highlight in job.get("highlights", []):
                    doc.add_paragraph(highlight, style='List Bullet')
                    
            # Skills
            doc.add_heading("SKILLS", level=1)
            skills = resume_data.get("skills", {})
            skills_str = ", ".join([f"{k.capitalize()}: {', '.join(v)}" for k, v in skills.items() if v])
            doc.add_paragraph(skills_str)
            
            # Education
            doc.add_heading("EDUCATION", level=1)
            for edu in resume_data.get("education", []):
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')}").bold = True
                p.add_run(f", {edu.get('institution', '')} ({edu.get('graduation_year', '')})")

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Failed to generate DOCX: {e}")
            return False
            
    @staticmethod
    def export_cover_letter(letter_text: str, personal_info: dict, company_name: str, output_path: str) -> bool:
        """Generates cover letter DOCX."""
        try:
            doc = Document()
            doc.add_heading(personal_info.get("full_name", ""), level=0)
            doc.add_paragraph(f"{personal_info.get('location', '')} | {personal_info.get('email', '')}")
            
            for paragraph in letter_text.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
                    
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Failed to generate cover letter DOCX: {e}")
            return False
